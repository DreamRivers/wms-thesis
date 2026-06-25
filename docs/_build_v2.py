"""
完整重建 build-thesis-docx.py
- 读取 02-毕业论文.md
- 按章/节/段落生成 Word 文档
- 应用学校 2026 模板规格
"""
import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\meng\Desktop\wms-graduation\docs\02-毕业论文_新版.docx"
SRC = r"C:\Users\meng\Desktop\wms-graduation\docs\02-毕业论文.md"

F_HEI = "黑体"
F_SONG = "宋体"
F_KAI = "楷体"
F_FANG = "仿宋_GB2312"
F_TIMES = "Times New Roman"

SCHOOL = "重庆文理学院"
DEPT = "数学与人工智能学院"
MAJOR = "计算机科学与技术"
TITLE_LINE1 = "基于 Spring Boot + Vue 3 的"
TITLE_LINE2 = "电商仓储物资管理系统的设计与实现"
STUDENT = "江童"
STUDENT_ID = "202258914282"
TEACHER = "任占广"
TEACHER_TITLE = "讲师"
FINISH_TIME = "2026 年 06 月"
SCHOOL_CODE = "10642"
SECRET_LEVEL = "公  开"

def set_cn_font(run, name=F_SONG, size=12, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    # 全部设为同一字体名(中文走 eastAsia,西文走 ascii/hAnsi)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)

def setup_heading_styles(doc):
    """Word 标题样式按学校 2026 模板
    H1 章: 黑体小三(15pt) + 居中 + 1.5 倍行距 + 段前 0 + 段后 11 磅
    H2 节: 黑体四号(14pt) + 居左 + 1.5 倍行距 + 段前 0.5 行 + 段后 0
    H3 子节: 黑体小四(12pt) + 居左 + 1.5 倍行距 + 段前 0.5 行 + 段后 0
    """
    # Heading 1 - 章
    h1 = doc.styles['Heading 1']
    h1.font.size = Pt(15)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(12)  # 段后 1 行 (1.5 倍行距下 1 行 ≈ 12pt)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    h1.paragraph_format.page_break_before = True
    rPr1 = h1.element.get_or_add_rPr()
    rFonts1 = rPr1.find(qn('w:rFonts'))
    if rFonts1 is None:
        rFonts1 = OxmlElement('w:rFonts')
        rPr1.append(rFonts1)
    rFonts1.set(qn('w:eastAsia'), F_HEI)
    rFonts1.set(qn('w:ascii'), F_HEI)
    rFonts1.set(qn('w:hAnsi'), F_HEI)

    # Heading 2 - 节
    h2 = doc.styles['Heading 2']
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(12)  # 0.5 行 ≈ 12pt (1.5 倍行距下 24pt 一行,半行=12pt)
    h2.paragraph_format.space_after = Pt(0)
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    rPr2 = h2.element.get_or_add_rPr()
    rFonts2 = rPr2.find(qn('w:rFonts'))
    if rFonts2 is None:
        rFonts2 = OxmlElement('w:rFonts')
        rPr2.append(rFonts2)
    rFonts2.set(qn('w:eastAsia'), F_HEI)
    rFonts2.set(qn('w:ascii'), F_HEI)
    rFonts2.set(qn('w:hAnsi'), F_HEI)

    # Heading 3 - 子节
    h3 = doc.styles['Heading 3']
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(12)  # 0.5 行
    h3.paragraph_format.space_after = Pt(0)
    h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    rPr3 = h3.element.get_or_add_rPr()
    rFonts3 = rPr3.find(qn('w:rFonts'))
    if rFonts3 is None:
        rFonts3 = OxmlElement('w:rFonts')
        rPr3.append(rFonts3)
    rFonts3.set(qn('w:eastAsia'), F_HEI)
    rFonts3.set(qn('w:ascii'), F_HEI)
    rFonts3.set(qn('w:hAnsi'), F_HEI)

def add_para(doc, text, size=12, name=F_SONG, indent=True, first_line=Cm(0.74),
             line_rule=WD_LINE_SPACING.ONE_POINT_FIVE, before=0, after=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = line_rule
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.first_line_indent = first_line if indent else Cm(0)
    p.alignment = align
    if text:
        run = p.add_run(text)
        set_cn_font(run, name=name, size=size)
    return p

def add_title_center(doc, text, size=16, name=F_HEI, bold=True, after=Pt(12)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = after
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_cn_font(run, name=name, size=size, bold=bold)
    return p

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)
    return p

def add_h1(doc, text, odd_page=True):
    """章标题: 黑体小三(15pt) 居中 + 奇数页起始 (新节 oddPage)
    偶数页起始时,Word 自动插入空白页"""
    # 在新节中加 H1
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.start_type = WD_SECTION.ODD_PAGE
    # 章标题 (用 Heading 1 样式,样式里已设小三 15pt)
    p = doc.add_paragraph(style='Heading 1')
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_cn_font(run, name=F_HEI, size=15, bold=True)
    return p

def add_placeholder_page(doc):
    """插入 '此页无内容' 占位段 (用于 Word 自动加的奇偶空白页)"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("（此页无内容）")
    set_cn_font(run, name=F_SONG, size=12)

def add_h2(doc, text):
    """节标题: 黑体四号(14pt) 居左 + 段前 0.5 行"""
    p = doc.add_paragraph(style='Heading 2')
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_cn_font(run, name=F_HEI, size=14, bold=True)
    return p

def add_h3(doc, text):
    """子节标题: 黑体小四(12pt) 居左 + 段前 0.5 行"""
    p = doc.add_paragraph(style='Heading 3')
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_cn_font(run, name=F_HEI, size=12, bold=True)
    return p

def add_h3(doc, text):
    p = doc.add_paragraph(style='Heading 3')
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_cn_font(run, name=F_HEI, size=14, bold=True)
    return p

# ===========================
# 主程序
# ===========================
print('Loading md...')
with open(SRC, 'r', encoding='utf-8') as f:
    md = f.read()

# 找各章起止行
chapters = []
lines = md.split('\n')
chapter_starts = []
for i, l in enumerate(lines):
    m = re.match(r'^##\s+第\s+(\d+)\s+章\s+(.+)$', l)
    if m:
        chapter_starts.append((i, int(m.group(1)), m.group(2)))

# 切分各章
sections = []
for idx, (start, num, title) in enumerate(chapter_starts):
    end = chapter_starts[idx+1][0] if idx+1 < len(chapter_starts) else len(lines)
    sections.append((num, title, lines[start:end]))

print(f'Chapters: {[(n,t) for n,t,_ in sections]}')

# 找参考文献、致谢、附录
specials = []
for i, l in enumerate(lines):
    if l.startswith('## 参考文献'):
        specials.append(('refs', i, l))
    elif l.startswith('## 致谢') or l.startswith('##  致谢') or l.startswith('## 致  谢'):
        specials.append(('thanks', i, l))
    elif l.startswith('## 附录') or l.startswith('##  附录') or l.startswith('## 附  录'):
        specials.append(('appendix', i, l))
    elif l.startswith('## 摘要'):
        specials.append(('abstract_cn', i, l))
    elif l.startswith('## Abstract'):
        specials.append(('abstract_en', i, l))
print(f'Specials: {[(k,i) for k,i,_ in specials]}')

# 提取摘要/致谢/参考文献内容
abstract_cn_text = ''
abstract_en_text = ''
for k, i, _ in specials:
    if k == 'abstract_cn':
        # 找下一个 # 标题 或空
        end = len(lines)
        for j in range(i+1, len(lines)):
            if lines[j].startswith('##') or lines[j].startswith('### 关键'):
                if j > i+1:
                    end = j
                    break
        abstract_cn_text = '\n'.join(lines[i+1:end]).strip()
    elif k == 'abstract_en':
        end = len(lines)
        for j in range(i+1, len(lines)):
            if lines[j].startswith('##') or lines[j].startswith('### Key'):
                if j > i+1:
                    end = j
                    break
        abstract_en_text = '\n'.join(lines[i+1:end]).strip()

# 提取正文
chapter_sections = []
for num, title, content_lines in sections:
    chapter_sections.append((num, title, content_lines))

# ===========================
# 生成 docx
# ===========================
print('Creating docx...')
doc = Document()
setup_heading_styles(doc)

# 页面设置
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.1)
section.left_margin = Cm(2.1)
section.right_margin = Cm(2.1)
section.gutter = Cm(1.0)  # 装订线 1.0cm
section.header_distance = Cm(2.5)  # 页眉 2.5cm
section.footer_distance = Cm(2.0)  # 页脚 2cm
# 启用 different_first_page (封面不编页码) - 不,封面其实没页码因为它根本不在正文 section
# 启用 different_odd_even_pages (奇偶不同)
doc.settings.odd_and_even_pages_header_footer = True
# 打开时自动更新所有域 (让 TOC 自动生成)
doc.settings.element.append(OxmlElement('w:updateFields'))
# updateFields 元素需要 w:val="true"
updateFields_elem = doc.settings.element.find(qn('w:updateFields'))
if updateFields_elem is not None:
    updateFields_elem.set(qn('w:val'), 'true')

# ========== 封面 ==========
# 顶部:单位代码 + 密级
p = doc.add_paragraph()
p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT)
run = p.add_run("单位代码  ")
set_cn_font(run, name=F_SONG, size=14)
run = p.add_run(f"{SCHOOL_CODE}                              ")
set_cn_font(run, name=F_SONG, size=14)
run = p.add_run(f"密    级    {SECRET_LEVEL}")
set_cn_font(run, name=F_SONG, size=14)

# 校徽
LOGO = r"C:\Users\meng\Desktop\wms-graduation\docs\_template_imgs\image2.png"
if os.path.exists(LOGO):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(LOGO, width=Cm(3))

add_para(doc, "", indent=False, before=0, after=0)

# 校训
MOTTO = r"C:\Users\meng\Desktop\wms-graduation\docs\_template_imgs\image1.png"
if os.path.exists(MOTTO):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(MOTTO, width=Cm(13))

add_para(doc, "", indent=False, before=0, after=0)

# 本科毕业设计
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run("本科毕业设计")
set_cn_font(run, name=F_HEI, size=36, bold=True)

add_para(doc, "", indent=False, before=0, after=0)

# 题目
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run("题 目:")
set_cn_font(run, name=F_HEI, size=22, bold=True)
run = p.add_run("  ")
set_cn_font(run, name=F_SONG, size=15)
run = p.add_run(TITLE_LINE1 + TITLE_LINE2)
set_cn_font(run, name=F_SONG, size=15)
run.font.underline = True

add_para(doc, "", indent=False, before=0, after=0)
add_para(doc, "", indent=False, before=0, after=0)
add_para(doc, "", indent=False, before=0, after=0)

# 学生信息
info_items = [
    ("学       院:", DEPT),
    ("专       业:", MAJOR),
    ("学  生  姓  名:", STUDENT),
    ("学       号:", STUDENT_ID),
]
for label, val in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(label)
    set_cn_font(run, name=F_HEI, size=15, bold=True)
    run = p.add_run("  ")
    set_cn_font(run, name=F_SONG, size=15)
    run = p.add_run(val)
    set_cn_font(run, name=F_SONG, size=15, bold=True)
    run.font.underline = True

# 校内
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.left_indent = Cm(4)
p.paragraph_format.space_after = Pt(8)
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run("校 内 导 师:")
set_cn_font(run, name=F_HEI, size=15, bold=True)
run = p.add_run("  姓  名  ")
set_cn_font(run, name=F_HEI, size=15, bold=True)
run = p.add_run(TEACHER)
set_cn_font(run, name=F_SONG, size=15, bold=True)
run.font.underline = True
run = p.add_run("  职  称  ")
set_cn_font(run, name=F_HEI, size=15, bold=True)
run = p.add_run(TEACHER_TITLE)
set_cn_font(run, name=F_SONG, size=15, bold=True)
run.font.underline = True

# 校外
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.left_indent = Cm(4)
p.paragraph_format.space_after = Pt(8)
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run("校 外 导 师:")
set_cn_font(run, name=F_HEI, size=15, bold=True)
run = p.add_run("  姓  名  ")
set_cn_font(run, name=F_HEI, size=15, bold=True)
run = p.add_run("        ")
set_cn_font(run, name=F_SONG, size=15, bold=True)
run.font.underline = True
run = p.add_run("  职  称  ")
set_cn_font(run, name=F_HEI, size=15, bold=True)
run = p.add_run("        ")
set_cn_font(run, name=F_SONG, size=15, bold=True)
run.font.underline = True

# 完成时间
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.left_indent = Cm(4)
p.paragraph_format.space_after = Pt(8)
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run("完  成  时  间:")
set_cn_font(run, name=F_HEI, size=15, bold=True)
run = p.add_run("  ")
set_cn_font(run, name=F_SONG, size=15)
run = p.add_run(FINISH_TIME)
set_cn_font(run, name=F_SONG, size=15, bold=True)
run.font.underline = True

# 封面后加新 section (让封面独立) - 不强制分页 (add_section 本身就是分页符)
# 用 CONTINUOUS 让分节符不强制分页
_doc_new_sec = doc.add_section(WD_SECTION.CONTINUOUS)
_doc_new_sec.start_type = WD_SECTION.CONTINUOUS

# 但 cover 末尾还是要分页到摘要
# 让 cover 段尾加 page break

# ========== 中文摘要 ==========
# 标题: 黑体 + 居中 + 三号(16pt) + 1.5 倍行距
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.first_line_indent = Cm(0)
title_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(0)
run = title_p.add_run("摘  要")
set_cn_font(run, name=F_HEI, size=16, bold=True)

if abstract_cn_text:
    # 摘要正文: 宋体小四 + 22 磅固定行距 + 段前段后 0 + 首行缩进 2 字符
    # 数字英文 Times New Roman
    paras = abstract_cn_text.split('\n\n')
    for p in paras:
        p = p.strip()
        if not p: continue
        # 关键词行检测
        if '关键词' in p and (';' in p or '；' in p):
            # 关键词行: 与摘要正文间隔一行 (段前 1 行 22pt)
            clean = p.replace('**关键词**:', '').replace('**关键词**:', '').replace('**关键词**:', '').replace('关键词:', '').replace('关键词：', '').replace('**', '').strip()
            p_obj = doc.add_paragraph()
            p_obj.paragraph_format.first_line_indent = Cm(0)
            p_obj.paragraph_format.space_after = Pt(0)
            p_obj.paragraph_format.space_before = Pt(22)  # 间隔一行
            p_obj.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            run = p_obj.add_run("关键词:")
            set_cn_font(run, name=F_HEI, size=12, bold=True)
            run = p_obj.add_run("  " + clean)
            set_cn_font(run, name=F_FANG, size=12)
        else:
            # 摘要正文: 宋体小四 + 22 磅固定行距 + 段前段后 0 + 首行缩进 2 字符
            p_obj = doc.add_paragraph()
            p_obj.paragraph_format.first_line_indent = Cm(0.74)  # 2 字符
            p_obj.paragraph_format.space_after = Pt(0)
            p_obj.paragraph_format.space_before = Pt(0)
            p_obj.paragraph_format.line_spacing = Pt(22)  # 22 磅固定值
            p_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            parts = re.split(r'(\*\*[^*]+\*\*)', p)
            for part in parts:
                if not part: continue
                if part.startswith('**') and part.endswith('**'):
                    run = p_obj.add_run(part[2:-2])
                    set_cn_font(run, name=F_SONG, size=12, bold=True)
                else:
                    run = p_obj.add_run(part)
                    set_cn_font(run, name=F_SONG, size=12)

add_page_break(doc)

# ========== Abstract ==========
# 标题: Times New Roman + 加粗 + 居中 + 三号(16pt) + 1.5 倍行距
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.first_line_indent = Cm(0)
title_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(0)
run = title_p.add_run("Abstract")
set_cn_font(run, name=F_TIMES, size=16, bold=True)

if abstract_en_text:
    # Abstract 正文: Times New Roman 小四 + 22 磅固定行距 + 段前段后 0 + 首行缩进 2 字符
    paras = abstract_en_text.split('\n\n')
    for p in paras:
        p = p.strip()
        if not p: continue
        if 'Key Words' in p and ';' in p:
            # Key Words: Times New Roman + 加粗 + 小四 + 间隔一行 + 英文分号
            clean = p.replace('**Key Words**:', '').replace('**Key Words**:', '').replace('Key Words:', '').replace('**', '').strip()
            p_obj = doc.add_paragraph()
            p_obj.paragraph_format.first_line_indent = Cm(0)
            p_obj.paragraph_format.space_before = Pt(22)  # 间隔一行
            p_obj.paragraph_format.space_after = Pt(0)
            p_obj.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            run = p_obj.add_run("Key Words: ")
            set_cn_font(run, name=F_TIMES, size=12, bold=True)
            run = p_obj.add_run(clean)
            set_cn_font(run, name=F_TIMES, size=12)
        else:
            # Abstract 正文: Times New Roman + 小四 + 22 磅固定行距
            p_obj = doc.add_paragraph()
            p_obj.paragraph_format.first_line_indent = Cm(0.74)
            p_obj.paragraph_format.space_after = Pt(0)
            p_obj.paragraph_format.space_before = Pt(0)
            p_obj.paragraph_format.line_spacing = Pt(22)  # 22 磅固定值
            p_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            parts = re.split(r'(\*\*[^*]+\*\*)', p)
            for part in parts:
                if not part: continue
                if part.startswith('**') and part.endswith('**'):
                    run = p_obj.add_run(part[2:-2])
                    set_cn_font(run, name=F_TIMES, size=12, bold=True)
                else:
                    run = p_obj.add_run(part)
                    set_cn_font(run, name=F_TIMES, size=12)

add_page_break(doc)

# ========== 目录 ==========
# 标题: 三号黑体居中,字与字之间空 2 个空格
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.first_line_indent = Cm(0)
title_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(0)
run = title_p.add_run("目\u3000\u3000录")  # 三号黑体居中,字与字之间空 2 个全角空格(\u3000)
set_cn_font(run, name=F_HEI, size=16, bold=True)

# 生成自动目录 (Word TOC 域代码)
p = doc.add_paragraph()
pf = p.paragraph_format
pf.space_after = Pt(0)
pf.first_line_indent = Cm(0)
# Word TOC 域代码 - 收集所有 Heading 1-3
fldChar_begin = OxmlElement("w:fldChar")
fldChar_begin.set(qn("w:fldCharType"), "begin")
instrText = OxmlElement("w:instrText")
instrText.set(qn("xml:space"), "preserve")
instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
fldChar_sep = OxmlElement("w:fldChar")
fldChar_sep.set(qn("w:fldCharType"), "separate")
fldChar_end = OxmlElement("w:fldChar")
fldChar_end.set(qn("w:fldCharType"), "end")
run = p.add_run()
set_cn_font(run, name=F_SONG, size=12)
run._element.append(fldChar_begin)
run._element.append(instrText)
run._element.append(fldChar_sep)
run._element.append(fldChar_end)

add_page_break(doc)

# ========== 各章正文 ==========
def add_md_para(doc, text, list_remaining=False):
    """添加段落, 处理加粗标记 + 参考文献引用 [N] 右上角标
    学校规范: 宋体小四 + 22 磅固定行距 + 段前段后 0 + 首行缩进 2 字符
    """
    # 先按 **bold** 切分,再在每个 part 内切 [N]
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = Pt(22)  # 22 磅固定值
    pf.first_line_indent = Cm(0.74)  # 2 字符
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for part in parts:
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_cn_font(run, name=F_SONG, size=12, bold=True)
        else:
            # 在 part 内切 [N] 引用
            sub_parts = re.split(r'(\[\d+(?:[-,]\d+)*\])', part)
            for sub in sub_parts:
                if not sub: continue
                if re.match(r'^\[\d+(?:[-,]\d+)*\]$', sub):
                    # 右上角标
                    run = p.add_run(sub)
                    set_cn_font(run, name=F_SONG, size=12)
                    run.font.superscript = True
                else:
                    run = p.add_run(sub)
                    set_cn_font(run, name=F_SONG, size=12)
    return p

def add_md_list(doc, text, idx=None):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(22)  # 22 磅固定值
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("(" + str(idx) + ") " + text if idx else "- " + text)
    set_cn_font(run, name=F_SONG, size=12)

# 处理各章
for num, title, content_lines in chapter_sections:
    # 先加章标题 (Heading 1)
    # 章标题用中文数字 (一/二/三/四/五)
    cn_nums = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
    cn_num = cn_nums[num] if 0 < num < len(cn_nums) else str(num)
    chapter_title = f'第{cn_num}章  {title}'
    add_h1(doc, chapter_title)
    # 跳过章标题本身 (会作为 H1 添加)
    i = 1
    # 第一行是 ## 第 N 章 Xxx - 跳过
    # 但 num/title 已经知道
    while i < len(content_lines):
        line = content_lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith('### '):
            add_h2(doc, line[4:].strip())
        elif line.startswith('#### '):
            h4_text = line[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(h4_text)
            set_cn_font(run, name=F_HEI, size=13, bold=True)
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            add_md_list(doc, text)
        elif line.startswith('---'):
            pass  # 分隔符,跳过
        elif line.startswith('## '):
            # 遇到新 ## 标题 (参考文献/致谢),跳出本章节
            break
        else:
            # 段落
            add_md_para(doc, line)
        i += 1

# ========== 参考文献 ==========
add_page_break(doc)
add_h1(doc, "参考文献")
# 简单插入 (如果有)
ref_content = ""
for k, i, _ in specials:
    if k == 'refs':
        end = len(lines)
        for j in range(i+1, len(lines)):
            if lines[j].startswith('##') and '致谢' not in lines[j] and '附录' not in lines[j]:
                end = j
                break
        ref_content = '\n'.join(lines[i+1:end]).strip()
        break
if ref_content:
    for line in ref_content.split('\n'):
        line = line.strip()
        if not line: continue
        # 参考文献行: 不用 superscript (保留方括号),正常段落
        if line.startswith('- '):
            text = line[2:].strip()
        else:
            text = line
        # 加段落 (宋体五号 10.5pt + 固定行距 22 磅 + 悬挂缩进 + 段后 0)
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.63)  # 悬挂缩进 (五号 0.63cm)
        p.paragraph_format.line_spacing = Pt(22)  # 22 磅固定值
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_cn_font(run, name=F_SONG, size=10.5)  # 五号 = 10.5pt

# ========== 致谢 ==========
add_page_break(doc)
add_h1(doc, "致  谢")
thank_content = ""
for k, i, _ in specials:
    if k == 'thanks':
        end = len(lines)
        for j in range(i+1, len(lines)):
            if lines[j].startswith('##'):
                end = j
                break
        thank_content = '\n'.join(lines[i+1:end]).strip()
        break
if thank_content:
    for para in thank_content.split('\n\n'):
        p = para.strip()
        if not p: continue
        add_para(doc, p, size=12, name=F_SONG, indent=True,
                 first_line=Cm(0.74), line_rule=WD_LINE_SPACING.ONE_POINT_FIVE, before=0, after=0)

# ========== 附录 ==========
add_page_break(doc)
add_h1(doc, "附  录")
app_content = ""
for k, i, _ in specials:
    if k == 'appendix':
        end = len(lines)
        for j in range(i+1, len(lines)):
            if lines[j].startswith('##'):
                end = j
                break
        app_content = '\n'.join(lines[i+1:end]).strip()
        break
if app_content:
    for para in app_content.split('\n\n'):
        p = para.strip()
        if not p: continue
        if p.startswith('### '):
            add_h3(doc, p[4:].strip())
        else:
             add_para(doc, p, size=12, name=F_SONG, indent=True,
                     first_line=Cm(0.74), line_rule=WD_LINE_SPACING.ONE_POINT_FIVE, before=0, after=0)

# ========== 页眉 + 页脚 ==========
# 文档 sections 布局:
# s0: 封面(无页码无页眉)
# s1: 摘要(罗马 I)
# s2: Abstract(罗马 II)
# s3: 目录(罗马 III)
# s4: 第1章 ... s11: 附录 (阿拉伯 1, 2, 3...)
TITLE_SHORT = "基于Spring Boot+Vue3的电商仓储管理系统"
HEADER_EVEN = "重庆文理学院2026届计算机科学与技术专业本科毕业设计"

def set_header_footer(section, page_num_format='decimal', start_at=None, has_header=True, no_page_number=False):
    """设置页眉(奇数: 题目;偶数: 学校) + 页脚(页码) - default + even 都要设
    no_page_number=True: 不加 PAGE 字段 (用于封面)"""
    def _add_page_field(footer_obj, fmt='decimal'):
        footer_obj.is_linked_to_previous = False
        # 清空旧内容
        for p in footer_obj.paragraphs:
            for r in p.runs:
                r.text = ''
        if not footer_obj.paragraphs:
            footer_obj.add_paragraph()
        fp = footer_obj.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.first_line_indent = Cm(0)
        # 加页码字段
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        if fmt == 'lowerRoman':
            instrText.text = ' PAGE \\* roman \\* MERGEFORMAT '
        else:
            instrText.text = ' PAGE \\* MERGEFORMAT '
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        cached_run = OxmlElement('w:r')
        cached_rPr = OxmlElement('w:rPr')
        cached_sz = OxmlElement('w:sz')
        cached_sz.set(qn('w:val'), '21')
        cached_rPr.append(cached_sz)
        cached_run.append(cached_rPr)
        cached_t = OxmlElement('w:t')
        cached_t.text = '1'
        cached_run.append(cached_t)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run = fp.add_run()
        set_cn_font(run, name=F_SONG, size=10.5)
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        run._element.append(cached_run)
        run._element.append(fldChar3)

    # default footer (奇数页) + even footer (偶数页) 都要 PAGE 字段
    if not no_page_number:
        _add_page_field(section.footer, page_num_format)
        _add_page_field(section.even_page_footer, page_num_format)

    # 设置页码格式
    sectPr = section._sectPr
    # 移除旧 pgNumType
    for old in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(old)
    pgNumType = OxmlElement('w:pgNumType')
    if page_num_format == 'lowerRoman':
        pgNumType.set(qn('w:fmt'), 'lowerRoman')
    else:
        pgNumType.set(qn('w:fmt'), 'decimal')
    if start_at is not None:
        pgNumType.set(qn('w:start'), str(start_at))
    sectPr.append(pgNumType)

    # 页眉 - 奇偶不同
    if has_header:
        # 奇数页页眉: 题目 (默认 header = 奇数页)
        h_odd = section.header
        h_odd.is_linked_to_previous = False
        if h_odd.paragraphs:
            for p in h_odd.paragraphs:
                p.clear()
        if not h_odd.paragraphs:
            h_odd.add_paragraph()
        hop = h_odd.paragraphs[0]
        hop.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hop.paragraph_format.first_line_indent = Cm(0)
        run = hop.add_run(TITLE_SHORT)
        set_cn_font(run, name=F_SONG, size=10.5)

        # 偶数页页眉: 学校名
        h_even = section.even_page_header
        h_even.is_linked_to_previous = False
        if h_even.paragraphs:
            for p in h_even.paragraphs:
                p.clear()
        if not h_even.paragraphs:
            h_even.add_paragraph()
        hep = h_even.paragraphs[0]
        hep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hep.paragraph_format.first_line_indent = Cm(0)
        run = hep.add_run(HEADER_EVEN)
        set_cn_font(run, name=F_SONG, size=10.5)
    else:
        # 清空页眉 (封面/摘要/Abstract/目录无页眉)
        for hf in [section.header, section.even_page_header, section.first_page_header]:
            hf.is_linked_to_previous = False
            if hf.paragraphs:
                for p in hf.paragraphs:
                    p.clear()
        # 页脚保留 (前面已设了 PAGE 字段,前置部分也要显示罗马页码)
        # 不清空
        # 但封面是 s0 第一页,不应该显示页码
        # 用 different_first_page + first_page_footer 空
        section.different_first_page_header_footer = True
        first_footer = section.first_page_footer
        first_footer.is_linked_to_previous = False
        if first_footer.paragraphs:
            for p in first_footer.paragraphs:
                p.clear()
        if not first_footer.paragraphs:
            first_footer.add_paragraph()

# 设置每节
# s0: 封面 - 无页码无页眉
# s1: 摘要 - 罗马数字 I
# s2: Abstract - 罗马 II
# s3: 目录 - 罗马 III
# s4-s11: 章 (奇数页起始 + 阿拉伯 1,2,3...)
# 实际上,sections 顺序取决于 add_section 的位置
# 当前: s0=封面 (default section), s1-s8 = 8 个章/参考文献/致谢/附录
# 但前面 (摘要/Abstract/目录) 都在 s0 里!
# 我需要重新规划 sections

# 简化做法: 对 s0 整段(摘要/Abstract/目录) 用罗马数字,不设页眉
# 后面 s1-s8 是章,设阿拉伯 + 页眉

if len(doc.sections) >= 10:
    # s0 = 封面 (无页码无页眉)
    set_header_footer(doc.sections[0], 'decimal', has_header=False, no_page_number=True)
    doc.sections[0].different_first_page_header_footer = False
    # 移除 s0 first footer ref (确保 first 不指向空 footer)
    s0_sectPr = doc.sections[0]._sectPr
    for fr in list(s0_sectPr.findall(qn('w:footerReference'))):
        if fr.get(qn('w:type')) == 'first':
            s0_sectPr.remove(fr)
    for hr in list(s0_sectPr.findall(qn('w:headerReference'))):
        if hr.get(qn('w:type')) == 'first':
            s0_sectPr.remove(hr)
    # s1 = 摘要/Abstract/目录 (罗马, 起始 I, 无页眉)
    set_header_footer(doc.sections[1], 'lowerRoman', start_at=1, has_header=False)
    doc.sections[1].different_first_page_header_footer = False
    # 移除 s1 first footer/header ref
    s1_sectPr = doc.sections[1]._sectPr
    for fr in list(s1_sectPr.findall(qn('w:footerReference'))):
        if fr.get(qn('w:type')) == 'first':
            s1_sectPr.remove(fr)
    for hr in list(s1_sectPr.findall(qn('w:headerReference'))):
        if hr.get(qn('w:type')) == 'first':
            s1_sectPr.remove(hr)
    # s2-s9 = 第1-7章 + 参考文献/致谢/附录 (阿拉伯)
    for i, sec in enumerate(doc.sections[2:], start=2):
        if i == 2:
            set_header_footer(sec, 'decimal', start_at=1, has_header=True)
        else:
            s2_sectPr = doc.sections[2]._sectPr
            cur_sectPr = sec._sectPr
            for fr in s2_sectPr.findall(qn('w:footerReference')):
                new_fr = OxmlElement('w:footerReference')
                new_fr.set(qn('w:type'), fr.get(qn('w:type')))
                new_fr.set(qn('r:id'), fr.get(qn('r:id')))
                cur_sectPr.append(new_fr)
            for hr in s2_sectPr.findall(qn('w:headerReference')):
                new_hr = OxmlElement('w:headerReference')
                new_hr.set(qn('w:type'), hr.get(qn('w:type')))
                new_hr.set(qn('r:id'), hr.get(qn('r:id')))
                cur_sectPr.append(new_hr)
            for old in cur_sectPr.findall(qn('w:pgNumType')):
                cur_sectPr.remove(old)
            pgNumType = OxmlElement('w:pgNumType')
            pgNumType.set(qn('w:fmt'), 'decimal')
            cur_sectPr.append(pgNumType)

# 保存
doc.save(OUT)
print(f'OK: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
