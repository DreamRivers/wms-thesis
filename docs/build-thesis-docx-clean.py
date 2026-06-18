#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""毕业论文 Word 文档 - 简洁稳定版 (按甘柳峡 PDF 排版)"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\meng\Desktop\wms-graduation\docs\02-毕业论文.docx"

# 字体
F_HEI = "黑体"
F_SONG = "宋体"
F_KAI = "楷体"
F_TIMES = "Times New Roman"

# 文档信息
SCHOOL = "重庆文理学院"
DEPT = "数学与人工智能学院"
MAJOR = "计算机科学与技术"
TITLE_LINE1 = "基于 Spring Boot + Vue 3 的"
TITLE_LINE2 = "电商仓储物资管理系统的设计与实现"
STUDENT = "江童"
STUDENT_ID = "202258914282"
TEACHER = "任占广"
TEACHER_TITLE = "讲师"
FINISH_TIME = "2026 年 06 月"
SCHOOL_CODE = "10642"
SECRET_LEVEL = "公  开"


def set_cn_font(run, name=F_SONG, size=12, bold=False):
    """中英文字体设置"""
    run.font.name = F_TIMES
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), F_TIMES)
    rFonts.set(qn('w:hAnsi'), F_TIMES)


def setup_heading_styles(doc):
    """让 Word 标题样式使用黑体 + 按学校 2026 模板规格
    标题 1 (章) - 22pt 黑体加粗居中
    标题 2 (节) - 16pt 黑体加粗左对齐
    标题 3 (小节) - 14pt 黑体加粗左对齐
    """
    # Heading 1 - 章标题 (22pt 黑体加粗居中)
    h1 = doc.styles['Heading 1']
    h1.font.name = F_HEI
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(18)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = True
    rPr1 = h1.element.get_or_add_rPr()
    rFonts1 = rPr1.find(qn('w:rFonts'))
    if rFonts1 is None:
        rFonts1 = OxmlElement('w:rFonts')
        rPr1.append(rFonts1)
    rFonts1.set(qn('w:eastAsia'), F_HEI)
    rFonts1.set(qn('w:ascii'), F_TIMES)
    rFonts1.set(qn('w:hAnsi'), F_TIMES)
    # Heading 2 - 节标题 (16pt 黑体加粗左对齐)
    h2 = doc.styles['Heading 2']
    h2.font.name = F_HEI
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    h2.paragraph_format.keep_with_next = True
    rPr2 = h2.element.get_or_add_rPr()
    rFonts2 = rPr2.find(qn('w:rFonts'))
    if rFonts2 is None:
        rFonts2 = OxmlElement('w:rFonts')
        rPr2.append(rFonts2)
    rFonts2.set(qn('w:eastAsia'), F_HEI)
    rFonts2.set(qn('w:ascii'), F_TIMES)
    rFonts2.set(qn('w:hAnsi'), F_TIMES)
    # Heading 3 - 小节标题 (14pt 黑体加粗左对齐)
    h3 = doc.styles['Heading 3']
    h3.font.name = F_HEI
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    h3.paragraph_format.keep_with_next = True
    rPr3 = h3.element.get_or_add_rPr()
    rFonts3 = rPr3.find(qn('w:rFonts'))
    if rFonts3 is None:
        rFonts3 = OxmlElement('w:rFonts')
        rPr3.append(rFonts3)
    rFonts3.set(qn('w:eastAsia'), F_HEI)
    rFonts3.set(qn('w:ascii'), F_TIMES)
    rFonts3.set(qn('w:hAnsi'), F_TIMES)


def add_para(doc, text="", size=12, name=F_SONG, bold=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True,
             first_line=Cm(0.74), line_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
             before=0, after=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = line_rule
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.first_line_indent = first_line if indent else Cm(0)
    p.alignment = align
    if text:
        run = p.add_run(text)
        set_cn_font(run, name=name, size=size, bold=bold)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)
    return p


def add_title_center(doc, text, size=16, name=F_HEI):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_cn_font(run, name=name, size=size, bold=True)
    return p


def add_h1(doc, text):
    """一级标题 - 用 Heading 1 样式(Word 导航栏可识别)"""
    p = doc.add_paragraph(style='Heading 1')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_cn_font(run, name=F_HEI, size=16, bold=True)
    return p


def add_h2(doc, text):
    """二级标题 - 用 Heading 2 样式"""
    p = doc.add_paragraph(style='Heading 2')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_cn_font(run, name=F_HEI, size=14, bold=True)
    return p


def add_h3(doc, text):
    """三级标题 - 用 Heading 3 样式"""
    p = doc.add_paragraph(style='Heading 3')
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_cn_font(run, name=F_HEI, size=13, bold=True)
    return p


def add_inline_para(doc, parts, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_line=Cm(0.74), before=0, after=0):
    """parts: [(text, {'name':..., 'size':..., 'bold':...}), ...]"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.first_line_indent = first_line if indent else Cm(0)
    p.alignment = align
    for text, opts in parts:
        run = p.add_run(text)
        set_cn_font(run,
                    name=opts.get('name', F_SONG),
                    size=opts.get('size', 12),
                    bold=opts.get('bold', False))
    return p


add_list_item(doc, t)

# ============================================
# 5.6 分项设计
# ============================================
add_h2(doc, "5.6  分项设计")
add_para(doc, "本节对系统核心模块进行分项详细设计,说明各模块的类结构、关键流程、核心算法与接口契约,为后续编码实现提供精确指导。")

add_h3(doc, "5.6.1  入库管理模块设计")
add_para(doc, "(1) 模块类结构。入库管理模块的核心类包括 InboundOrder(实体)、InboundOrderItem(明细实体)、InboundController(控制器)、InboundOrderService(业务接口)、InboundOrderServiceImpl(业务实现)与 InboundOrderItemMapper(明细 Mapper)。其中,主单与明细采用一对多关联,通过 orderId 关联;Service 层封装业务规则,Controller 层仅负责请求参数解析与响应封装。")
add_para(doc, "(2) 关键流程。入库单的状态机为 DRAFT → PENDING → APPROVED → EXECUTING → FINISHED(或 CANCELED)。状态转移规则:DRAFT → PENDING 由用户点击提交触发;PENDING → APPROVED 由两级审核通过触发;APPROVED → EXECUTING 由执行时更新实际数量触发;EXECUTING → FINISHED 由完成后写入库存触发;任意非 FINISHED 状态可转为 CANCELED。")
add_para(doc, "(3) 关键算法。库存写入采用 SELECT ... FOR UPDATE 行级锁:在事务内先按 (goodsId, locationId, batchNo) 锁定库存记录,若不存在则创建并初始化为 0,再 quantity = quantity + qty 更新;同时 batchInsert 库存流水表,保证并发安全与数据一致性。")
add_para(doc, "(4) 接口契约。入库模块提供以下核心接口:POST /inbound/order/save 保存草稿;POST /inbound/order/submit/{id} 提交草稿;POST /inbound/order/audit 审核;POST /inbound/order/execute/{id} 执行;POST /inbound/order/complete/{id} 完成入库;POST /inbound/order/cancel/{id} 作废入库单;DELETE /inbound/order/{id} 删除草稿(仅 DRAFT 状态)。")

add_h3(doc, "5.6.2  出库管理模块设计")
add_para(doc, "出库模块与入库模块结构对称,采用两级审批流程:部门负责人审批 → 仓管人员审批 → 发货 → 完成。核心差异在于出库涉及库存扣减而非增加,需严格校验库存是否充足。")
add_para(doc, "(1) 关键算法。发货时,事务内对每条明细执行:1) SELECT ... FOR UPDATE 锁定库存记录;2) 校验 quantity >= qty,不足则抛出 BizException(STOCK_NOT_ENOUGH);3) quantity = quantity - qty、availableQty = quantity - lockedQty;4) 写入 wms_stock_record 流水,变更类型为 OUTBOUND。")
add_para(doc, "(2) 状态机。APPLY → APPROVING(step=1 部门) → APPROVED(step=2 仓管) → SHIPPED → FINISHED,或任意阶段 → CANCELED。")

add_h3(doc, "5.6.3  库存与盘点模块设计")
add_para(doc, "(1) 库存查询。库存列表支持按 goodsId、warehouseId、locationId、batchNo 多条件过滤,采用 MyBatis-Plus 动态条件构造器;分页采用 MyBatis-Plus Page<T> 封装;响应包含当前数量 quantity、锁定数量 lockedQty、可用数量 availableQty(即 quantity - lockedQty)。")
add_para(doc, "(2) 盘点流程。盘点单状态机:DRAFT → RECORDED → ADJUSTED。创建盘点单时,系统自动拉取当前所有 (goodsId, locationId, batchNo) 的账面数量 systemQty;盘点员录入 actualQty;调整阶段根据差值 diff = actualQty - systemQty 自动修正库存,并写盘亏/盘盈流水。")

add_h3(doc, "5.6.4  报表统计模块设计")
add_para(doc, "报表模块基于 MySQL 聚合查询实现,核心接口包括:GET /report/dashboard 返回今日入库数、出库数、库存预警数、待审核数等汇总指标;GET /report/inbound/trend?days=7 返回近 N 天入库单数与数量趋势;GET /report/outbound/trend?days=7 返回出库趋势;GET /report/inventory/value 返回按商品类型统计的库存总价值;GET /report/topGoods?limit=10 返回 TOP10 出入库商品。聚合查询使用 GROUP BY DATE(createTime) 按天分组,补全无数据日期为 0,保证趋势图 X 轴连续。")

add_h3(doc, "5.6.5  用户与权限模块设计")
add_para(doc, "(1) 鉴权流程。1) 用户登录:POST /auth/login 接收 {username, password},UserDetailsService 加载用户信息,BCrypt 校验密码;2) 密码正确后,Sa-Token 调用 StpUtil.login(uid) 创建 token,写入 Redis(有效期 8 小时);3) 前端将 token 存入 localStorage,后续请求通过 Authorization Header 传递;4) 后端 Sa-Token 拦截器 SaInterceptor 自动校验 token 有效性,失效返回 401;5) 注解式权限:@SaCheckPermission(\"user:add\") 校验当前用户是否拥有该权限。")
add_para(doc, "(2) 动态路由。前端 Vue Router 在 router.beforeEach 钩子中:白名单(/login、/404)放行;未登录跳 /login;已登录从 Pinia store 获取 roles 与 permissionCodes,过滤 asyncRoutes 生成实际可访问路由表;每次路由跳转调用 auth/routes 接口(后端返回该用户的菜单树),保证前后端权限一致。")

add_h3(doc, "5.6.6  操作日志模块设计")
add_para(doc, "操作日志通过自定义 @Log 注解 + AOP 切面实现无侵入埋点。@Log(module=\"入库管理\", action=\"保存草稿\") 标注在 Controller 方法上,OperationLogAspect 拦截:正常返回记录 status=1、costTime、requestParam、response;异常记录 status=0、errorMsg、异常堆栈;写入 sys_operation_log 表,user_id 从 Sa-Token 当前用户获取。日志写入采用 @Async 异步执行,避免影响主业务性能;模块化设计支持后续扩展到 module 字典。")

# ============================================
# 第 6 章
# ============================================
    "第 5 章 系统实现:阐述开发环境搭建、公共模块实现、各功能模块的核心实现与关键技术。",
    "第 6 章 系统测试:介绍测试环境,进行功能测试、性能测试与兼容性测试。",
    "第 7 章 总结与展望:总结主要工作与成果,分析不足并提出改进方向。",
], 1):
    add_list_item(doc, t, idx=i)

add_page_break(doc)

# ============================================
# 第 2 章 设计理论与技术
# ============================================
add_h1(doc, "2  设计理论与技术")
add_para(doc, "本章阐述系统设计所涉及的设计理论基础与本系统所采用的关键技术。设计理论涵盖软件架构、数据库设计、接口设计、权限模型、前后端分离、缓存与并发控制等核心理论;技术部分(下一章)介绍本系统采用的具体技术框架与组件,为后续章节的设计与实现提供理论依据与技术支撑。")

add_h2(doc, "2.1  软件架构设计理论")
add_para(doc, "软件架构是软件系统的高层结构,规定了系统由哪些子系统构成、子系统之间的交互关系以及指导原则。常见的 Web 应用架构模式包括以下几类:")
add_para(doc, "(1) B/S 架构(Browser/Server)。B/S 架构将系统功能实现的核心部分集中到服务器上,客户端通过浏览器访问应用。该架构的优势在于部署与维护便捷、跨平台能力强,用户无需安装专门客户端即可使用。本系统采用 B/S 架构,前端为基于 Vue 3 构建的单页应用(SPA),通过浏览器访问部署在服务器上的后端服务。")
add_para(doc, "(2) MVC 模式(Model-View-Controller)。MVC 模式将应用划分为模型(数据与业务逻辑)、视图(用户界面)、控制器(请求处理与流程调度)三个部分,实现职责分离与高内聚低耦合。本系统后端采用 Spring MVC 框架,前端采用 Vue 3 的组合式 API 与 Pinia 状态管理,在不同层面体现了 MVC 思想。")
add_para(doc, "(3) 分层架构(Layered Architecture)。分层架构将系统划分为表示层、业务逻辑层、数据访问层与数据持久层,层与层之间通过接口调用,不允许跨层访问。本系统后端遵循 Controller→Service→Mapper→Database 的分层结构,各层职责清晰,便于维护与扩展。")

add_h2(doc, "2.2  数据库设计理论")
add_para(doc, "数据库设计是信息系统开发的核心环节,直接影响系统的数据一致性、查询效率与可维护性。常用的设计理论与方法包括:")
add_para(doc, "(1) 关系数据库范式理论。关系数据库的范式(Normal Form)是衡量关系模式规范化程度的标准。第三范式(3NF)要求关系模式在满足第一范式(原子性)与第二范式(部分函数依赖消除)的基础上,消除非主属性对主键的传递函数依赖。本系统主要业务表设计符合 3NF,例如:wms_inbound_order 存储单据主信息,wms_inbound_order_item 存储明细,实现主-明细分离;wms_stock 与 wms_stock_record 分离,前者记录当前库存,后者记录所有库存变动流水。")
add_para(doc, "(2) ER 图(Entity-Relationship Diagram)。ER 图通过实体、属性、联系三个要素描述现实世界的概念模型。本系统核心实体包括入库单、出库单、商品、库位、库存、库存流水、用户、角色、菜单等,实体间通过 1:N 或 N:N 关系建立联系。")
add_para(doc, "(3) 主键与外键设计。本系统采用数据库自增 BIGINT 作为主键,保证唯一性与查询效率;通过外键约束维护表间参照完整性。考虑到高并发场景下外键约束可能成为性能瓶颈,本系统对核心业务表(如 wms_stock)仅在逻辑层维护外键关系,不显式声明物理外键,以提升写入性能。")

add_h2(doc, "2.3  接口设计理论")
add_para(doc, "(1) RESTful 架构风格。REST(Representational State Transfer)是一种基于 HTTP 协议的分布式系统架构风格,强调以资源为中心,通过统一接口对资源进行 CRUD 操作。本系统接口设计遵循 RESTful 原则:使用 HTTP 动词表示操作类型(GET 查询、POST 创建、PUT 更新、DELETE 删除);URI 表示资源,采用名词复数形式;通过 HTTP 状态码表示请求结果;请求与响应采用 JSON 格式,字段命名采用驼峰式。")
add_para(doc, "(2) 统一响应封装。为便于前端统一处理,本系统定义统一响应格式 Result<T>,包含 code(业务状态码)、message(提示信息)、data(业务数据)三个字段。")

add_h2(doc, "2.4  权限模型设计理论")
add_para(doc, "(1) RBAC 模型(Role-Based Access Control)。RBAC 是基于角色的访问控制模型,通过用户-角色-权限三层关系实现权限管理。本系统采用 RBAC0 模型,核心表包括 sys_user(用户)、sys_role(角色)、sys_menu(权限资源)、sys_user_role 与 sys_role_menu(用户-角色、角色-权限的多对多关系)。")
add_para(doc, "(2) Sa-Token 鉴权框架。Sa-Token 是国产轻量级 Java 鉴权框架,提供登录认证、权限校验、会话管理、单点登录等功能。本系统采用 Sa-Token 的注解式鉴权(@SaCheckPermission、@SaCheckRole)与拦截器鉴权相结合的方式,实现接口级别的细粒度权限控制。")

add_h2(doc, "2.5  前后端分离架构理论")
add_para(doc, "前后端分离架构将系统的表示层与业务逻辑层解耦,前端专注于用户界面与交互体验,后端专注于业务处理与数据持久化,通过标准化的 API 接口(如 RESTful)进行数据交互。其优势包括:职责清晰(前后端开发团队可并行工作);独立部署(前端可作为静态资源部署到 CDN);技术异构(后端可替换为其他语言实现而不影响前端);可维护性高。本系统采用前后端分离架构,后端基于 Spring Boot 暴露 RESTful API,前端基于 Vue 3 + Vite 构建单页应用。")

add_h2(doc, "2.6  缓存与并发控制理论")
add_para(doc, "(1) 缓存理论基础。缓存是通过将热点数据存储在访问速度更快的存储介质中,以减少对后端存储的直接访问,提高系统响应速度。本系统采用 Redis 作为分布式缓存,主要缓存用户会话信息(Sa-Token token 存储)与业务数据字典。")
add_para(doc, "(2) 数据库并发控制。本系统在库存扣减等关键业务场景中,采用悲观锁(SELECT ... FOR UPDATE 行级锁)、事务隔离(REPEATABLE READ)、乐观锁(版本号字段)三种并发控制策略。")

add_h2(doc, "2.7  软件开发方法论")
add_para(doc, "本系统在开发过程中采用迭代式开发方法,将系统开发划分为需求分析、系统设计、系统实现、系统测试、部署上线等阶段,每个阶段产出明确的可交付物。在编码规范上遵循阿里巴巴 Java 开发手册与 Vue 官方风格指南,使用 Git 进行版本管理,采用语义化版本号(SemVer)管理发布版本。")

# ============================================
# 第 3 章 相关技术与方法 (原第 2 章)
# ============================================
add_h1(doc, "3  相关技术与方法")
add_para(doc, "本章介绍了本项目所用到的主要技术,选择 Java 进行后端开发,前端使用 Vue 3,采用前后端分离架构开发,使用 MySQL 和 Redis 进行数据存储。这些技术都较为成熟且相互之间具有良好的兼容性,有利于前后端分离开发,提高开发效率的同时保证系统的可维护性。")

add_h2(doc, "4.1  Spring Boot 框架")
add_para(doc, "Spring Boot 是 Pivotal 团队在 Spring 4.0 基础上推出的快速开发框架,采用&#8220;约定优于配置&#8221;理念,通过自动配置、起步依赖等特点,去除了传统 Spring 应用程序中复杂的 XML 配置,使开发者可以迅速地构造出独立运行的 Java 应用。")
add_para(doc, "本系统选用 Spring Boot 2.7.18,具有以下优势:")
for t in [
    "内嵌 Web 容器:内嵌 Tomcat,无需独立部署 WAR,简化部署过程;",
    "自动配置:根据项目的依存关系自动进行相应配置,极大提高开发效率;",
    "起步依赖:通过 Starter 依赖管理,快速集成 MyBatis、Druid 数据源、Hutool 工具集等组件;",
    "AOP 与日志:内置 spring-boot-starter-aop 配合自定义 @Log 注解,实现操作日志无侵入式记录;",
    "生态完善:完善的生态与社区支持,可与各种主流技术栈无缝整合。",
]:
    add_list_item(doc, t)

add_h2(doc, "4.2  MyBatis-Plus")
add_para(doc, "MyBatis-Plus(简称 MP)是 MyBatis 的增强工具,遵循&#8220;只做增强不做改变&#8221;的理念,在不改变 MyBatis 原有功能的基础上,简化开发、提高效率。本系统使用的核心特性包括:")
for t in [
    "通用 CRUD:BaseMapper 提供通用增删改查,免写 XML SQL 语句;",
    "链式查询:LambdaQueryWrapper 支持链式查询,字段引用避免硬编码字符串;",
    "分页插件:PaginationInnerInterceptor 内置分页插件,高效处理大数据量查询;",
    "逻辑删除:通过 application.yml 全局配置 logic-delete-field: deleted(配合 logic-delete-value: 1、logic-not-delete-value: 0)实现软删除,保留数据便于审计;",
    "自动填充:MetaObjectHandler 自动填充 create_time、update_time、create_by、update_by 四个公共字段;",
    "乐观锁:OptimisticLockerInnerInterceptor 通过 @Version 字段防止并发更新覆盖。",
]:
    add_list_item(doc, t)

add_h2(doc, "4.3  Sa-Token 鉴权框架")
add_para(doc, "Sa-Token 是一款轻量级 Java 权限认证框架,提供登录认证、权限认证、踢人下线等功能。本系统使用其核心特性:")
for t in [
    "基于 Token:生成 UUID 作为 Token,存储于 Redis,支持分布式 Session;",
    "RBAC 鉴权:内置基于角色的访问控制,通过 @SaCheckPermission 注解实现权限校验;",
    "注解鉴权:@SaIgnore 忽略鉴权、@SaCheckLogin 登录校验、@SaCheckPermission 权限校验(本项目实际仅使用 @SaIgnore 与登录态校验,权限拦截见 4.5 节);",
    "深度集成:与 Spring Boot 深度集成,配置简单,使用方便。",
]:
    add_list_item(doc, t)

add_h2(doc, "4.4  Vue 3 与 Element Plus")
add_para(doc, "Vue 3 是尤雨溪团队发布的渐进式 JavaScript 框架,采用 Composition API + Proxy 重写,性能大幅提升。Element Plus 是基于 Vue 3 的组件库,提供完整的中后台 UI 组件。本系统使用其核心组件:el-form 实现数据录入、el-input 数据输入、el-select 下拉选择、el-table 表格展示、el-dialog 弹窗、el-pagination 分页、el-tag 状态标签等。本项目还引入了以下关键前端技术:")
for t in [
    "Vue Router 4.3:实现 SPA 单页应用的多路由管理,根据用户角色动态生成菜单和路由;",
    "Pinia 2.1:Vue 3 官方推荐的状态管理库,集中存储用户信息、Token、路由权限等全局状态;",
    "Axios 1.6:HTTP 请求库,封装统一请求拦截器,自动携带 Authorization Token,统一处理响应与错误;",
    "ECharts 5.5 + vue-echarts 7:数据可视化,用于仪表盘折线图、饼图、库存趋势图等报表;",
    "Vite 5:新一代前端构建工具,支持 ES Module 热更新,开发与构建速度优于 Webpack。",
]:
    add_list_item(doc, t)

add_h2(doc, "4.5  MySQL 与 Redis")
add_para(doc, "MySQL 8.0 提供窗口函数、CTE 等高级特性,InnoDB 存储引擎支持行级锁与事务,确保数据一致性。本系统使用 utf8mb4_unicode_ci 字符集,支持 emoji 和特殊字符;通过联合唯一索引保证数据唯一性。")
add_para(doc, "Redis 作为缓存与 Session 存储,提供毫秒级响应:存储 Sa-Token 生成的 Token(由 sa-token-redis-jackson 集成)、缓存热点数据(如字典、配置项),减少数据库访问压力。Lettuce 作为默认 Redis 客户端,基于 Netty 实现 NIO 异步通信,支持连接池与高并发场景。")

add_h2(doc, "4.6  前后端分离架构")
add_para(doc, "本系统采用前后端分离架构,后端专注业务与数据,前端专注视图与交互,通过 JSON 格式的 RESTful API 进行通信。后端基于 Spring Boot 部署在 8080 端口,前端基于 Vite 构建后由 Nginx 或静态服务器部署,开发期通过 Vite Proxy(/api 前缀)转发到后端。该架构的优势包括:职责清晰,前后端职责明确,团队协作更高效;独立部署,前端可独立部署为静态资源(CDN 加速访问);横向扩展,后端可独立横向扩展,支持高并发访问;技术选型解耦,后端可替换为其他语言实现而不影响前端。")

add_page_break(doc)

# ============================================
# 第 3 章
# ============================================
add_h1(doc, "4  系统需求分析")

add_h2(doc, "5.1  可行性分析")
add_h3(doc, "4.1.1  技术可行性分析")
add_para(doc, "利用 Spring Boot + Vue 3 前后端分离技术进行管理系统的开发,在技术层面具有明显的优越性:Spring Boot 通过自动配置、起步依赖等功能,大大简化了工程建设与配置过程,内置 Servlet 容器,简化了部署过程;Vue 3 采用 Composition API + Proxy 带来性能提升,组件化设计使代码更易维护;IDEA、VS Code、Maven、Node.js 等开发工具完善,文档丰富,社区活跃。")

add_h3(doc, "5.1.2  经济可行性分析")
add_para(doc, "利用 Spring Boot + Vue 3 开发管理系统,从经济学角度来说是可行的:Spring Boot、Vue、MySQL、Redis 均为开源框架,无需支付许可费;拥有丰富的开源生态,提供成熟的组件与解决方案,降低开发工作量;模块化设计,可根据需要灵活扩展,避免大规模重构,降低后续维护成本。")

add_h3(doc, "5.1.3  操作可行性分析")
add_para(doc, "以 Spring Boot + Vue 3 为基础,建立了一个高可操作性的管理系统:Element Plus 组件化设计,符合用户操作习惯,减少学习成本;业务流程设计合理,操作提示完整,用户体验良好;项目目录与配置标准化,新开发者能够很快了解项目架构。")

add_h2(doc, "5.2  系统功能需求分析")
add_h3(doc, "4.2.1  用户角色")
add_para(doc, "本系统支持 4 类角色协同工作,各角色职责如下:")
add_three_line_table(doc,
    ["角色", "角色编码", "职责"],
    [
        ["系统管理员", "ADMIN", "用户管理、角色管理、菜单管理、系统配置"],
        ["仓库管理员", "WAREHOUSE", "商品/库位/供应商维护,入库/出库/盘点执行,库存管理"],
        ["部门负责人", "DEPT_LEADER", "审批本部门员工的出库申请"],
        ["普通员工", "EMPLOYEE", "提交出库申请,查询个人记录"]
    ],
    col_widths_cm=[3, 3, 9.5],
    caption="表 3.1  系统角色与职责"
)

add_h3(doc, "5.2.2  核心功能需求")
add_inline_para(doc, [
    ("(1) 用户管理:", {'name': F_HEI, 'size': 12, 'bold': True}),
    ("用户增删改查、角色分配、启停账号、密码重置、登录日志。", {'name': F_SONG, 'size': 12})
])
add_inline_para(doc, [
    ("(2) 商品管理:", {'name': F_HEI, 'size': 12, 'bold': True}),
    ("商品 SKU 信息维护(编码、名称、规格、单位、价格);商品分类管理,支持多级分类;安全库存、临期预警天数设置。", {'name': F_SONG, 'size': 12})
])
add_inline_para(doc, [
    ("(3) 库位管理:", {'name': F_HEI, 'size': 12, 'bold': True}),
    ("仓库信息维护;库位四级结构:仓库→库区→货架→库位;库位容量与状态管理。", {'name': F_SONG, 'size': 12})
])
add_inline_para(doc, [
    ("(4) 供应商管理:", {'name': F_HEI, 'size': 12, 'bold': True}),
    ("供应商信息维护(编码、名称、联系人、电话);信用管理。", {'name': F_SONG, 'size': 12})
])
add_inline_para(doc, [
    ("(5) 入库管理:", {'name': F_HEI, 'size': 12, 'bold': True}),
    ("入库单类型:采购入库、退货入库、调拨入库;状态机:草稿→待审→已审→执行中→已完成 / 已驳回 / 已作废;支持批次管理、保质期管理。", {'name': F_SONG, 'size': 12})
])
add_inline_para(doc, [
    ("(6) 出库管理:", {'name': F_HEI, 'size': 12, 'bold': True}),
    ("出库单类型:销售出库、领用出库、调拨出库、报损出库;两级审批流:部门负责人审核→仓库管理员审核;状态机:申请→审批中→已审批→拣货中→已发货→已完成 / 已驳回 / 已作废。", {'name': F_SONG, 'size': 12})
])
add_inline_para(doc, [
    ("(7) 库存管理:", {'name': F_HEI, 'size': 12, 'bold': True}),
    ("实时库存查询(按商品、库位、批次);库存流水记录,支持追溯;联合唯一索引 (goods_id, location_id, batch_no) 保证库存唯一性。", {'name': F_SONG, 'size': 12})
])
add_inline_para(doc, [
    ("(8) 盘点管理:", {'name': F_HEI, 'size': 12, 'bold': True}),
    ("盘点单类型:全盘、抽盘、动态盘点;自动抓取当前库存,录入实盘数;差异调整,写盘点调整流水。", {'name': F_SONG, 'size': 12})
])
add_inline_para(doc, [
    ("(9) 预警管理:", {'name': F_HEI, 'size': 12, 'bold': True}),
    ("低库存预警(定时扫描商品库存,低于安全库存时生成预警通知);临期预警(定时扫描接近保质期的商品);预警通知推送给相关角色。", {'name': F_SONG, 'size': 12})
])
add_inline_para(doc, [
    ("(10) 报表统计:", {'name': F_HEI, 'size': 12, 'bold': True}),
    ("仪表盘(总库存量、今日入库/出库、预警数量);出入库趋势图(近 7 天);低库存/临期预警列表。", {'name': F_SONG, 'size': 12})
])

add_h2(doc, "5.3  系统非功能需求")
for t in [
    "性能需求:单页面加载 ≤ 3 秒,接口响应 ≤ 1 秒;",
    "安全需求:BCrypt 密码加密,Sa-Token Token 鉴权,MyBatis-Plus 参数绑定防 SQL 注入;",
    "可用性需求:界面友好,操作提示完整,错误信息友好;",
    "可维护性需求:模块化设计,关键 Service / Controller 类均编写 JavaDoc 注释;",
    "可扩展性需求:系统架构支持功能扩展和二次开发。",
]:
    add_list_item(doc, t)

add_h2(doc, "5.4  业务流程分析")
add_h3(doc, "4.4.1  入库流程")
add_para(doc, "仓管创建草稿 → 提交待审 → 管理员审核(通过/驳回) → 仓管执行(逐行实收数量) → 完成(自动写库存 + 写库存流水)。")
add_para(doc, "关键点:草稿状态可编辑,提交后不可修改;审核通过后才可执行;执行时记录实际收货数量,可能与计划数量不一致;执行完成后自动更新 wms_stock 库存表,并写入 wms_stock_record 流水表。")

add_h3(doc, "5.4.2  出库流程")
add_para(doc, "员工申请(APPLY) → 部门负责人审核(step=1, 通过→APPROVING / 驳回→REJECTED) → 仓库管理员审核(step=2, 通过→APPROVED / 驳回→REJECTED) → 拣货发货(SHIPPED,行级锁扣减库存 + 写库存流水) → 完成(FINISHED,更新状态与完成时间)。")
add_para(doc, "关键点:两级审批均通过后才可发货,任一级驳回即终止;审批记录写入 wms_outbound_approval 表,便于审计;拣货发货时通过 SELECT ... FOR UPDATE 行级锁扣减库存,校验库存是否充足,不足时抛 STOCK_NOT_ENOUGH 业务异常,并写入 wms_stock_record 流水表。")

add_page_break(doc)

# ============================================
# 第 4 章
# ============================================
add_h1(doc, "5  系统设计")

add_h2(doc, "6.1  系统总体架构")
add_para(doc, "采用经典三层架构 + 前后端分离架构。浏览器通过 HTTP/JSON 与 Vue 3 前端通信,Vue 3 前端通过 /api/* 接口调用 Spring Boot 后端,后端通过 Controller、Service、Mapper、Sa-Token 等模块实现业务逻辑,数据存储于 MySQL 8 与 Redis 7。")
add_para(doc, "后端分包结构:com.wms 根包下包含 common(通用类 Result/异常/枚举/分页)、config(配置类 SaToken/Redis/MybatisPlus/Swagger)、framework(框架 注解/AOP/定时任务)以及 modules(业务模块)六大子系统。modules 进一步细分为 system(系统管理)、basic(主数据)、inbound(入库管理)、outbound(出库管理)、stock(库存管理)、report(报表统计)。")

add_h2(doc, "6.2  功能模块设计")
add_para(doc, "用户端功能模块包括:系统管理(用户/角色/菜单/日志)、主数据管理(商品/分类/仓库/库位/供应商)、入库管理(入库单列表/新增/详情)、出库管理(出库申请/我的申请/详情)、库存管理(库存查询/流水/预警)、盘点管理(盘点单/执行)、报表统计(仪表盘/业务报表)。")

add_h2(doc, "6.3  数据库设计")
add_h3(doc, "5.3.1  E-R 图(简化)")
add_para(doc, "系统 E-R 图描述实体间关系:用户-角色-菜单构成 RBAC 权限模型,商品-入库明细-入库单-供应商构成采购入库链路,仓库-库位-库存-库存流水构成库存管理链路,出库单-出库明细-审批流构成两级审批链路。各表均设有 create_time、update_time、create_by、update_by、deleted 等公共字段,符合企业级数据治理规范。")

add_h3(doc, "6.3.2  核心表结构")
add_para(doc, "系统共设计 21 张业务表(系统模块 6 张 + 主数据 5 张 + 入库模块 2 张 + 出库模块 3 张 + 库存模块 4 张 + 通知表 1 张),涵盖了从权限、主数据到单据、库存的全链路数据需求。下表列出 16 类核心数据表(含主表/明细的合并表示):")
add_three_line_table(doc,
    ["序号", "表名", "所属模块", "说明"],
    [
        ["1", "sys_user", "系统", "系统用户表"],
        ["2", "sys_role", "系统", "系统角色表"],
        ["3", "sys_menu", "系统", "菜单权限表"],
        ["4", "sys_user_role", "系统", "用户角色关联"],
        ["5", "sys_role_menu", "系统", "角色菜单关联"],
        ["6", "sys_operation_log", "系统", "操作日志"],
        ["7", "wms_category", "主数据", "商品分类"],
        ["8", "wms_warehouse", "主数据", "仓库"],
        ["9", "wms_location", "主数据", "库位"],
        ["10", "wms_supplier", "主数据", "供应商"],
        ["11", "wms_goods", "主数据", "商品"],
        ["12", "wms_inbound_order/item", "入库", "入库单主表/明细(2张)"],
        ["13", "wms_outbound_order/item/approval", "出库", "出库单主表/明细/审批(3张)"],
        ["14", "wms_stock/record", "库存", "实时库存/库存流水(2张)"],
        ["15", "wms_stocktaking_order/item", "盘点", "盘点单主表/明细(2张)"],
        ["16", "wms_notification", "通知", "系统通知表"],
    ],
    col_widths_cm=[1.2, 5, 1.8, 7],
    caption="表 4.1  系统业务数据表"
)
add_para(doc, "关键设计:wms_stock 联合唯一索引 (goods_id, location_id, batch_no),保证库存唯一性;主要业务表包含 create_time/update_time/create_by/update_by/deleted 等公共字段,便于审计与回收站;单据均采用&#8220;主表+明细表&#8221;双表结构,便于管理与追溯;wms_outbound_approval 审批流留痕,便于审计与回溯。")

add_h2(doc, "6.4  接口设计")
add_para(doc, "RESTful 风格,统一返回 Result<T> 格式:{ code: 200, message: &#8220;操作成功&#8221;, data: {...}, timestamp: ... }。")
add_inline_para(doc, [
    ("错误码:", {'name': F_HEI, 'size': 12, 'bold': True}),
    ("200 成功、400 参数错误、401 未登录、403 无权限、500 业务异常、1001 库存不足、1002 状态非法。", {'name': F_SONG, 'size': 12})
])
add_para(doc, "主要接口:POST /auth/login 用户登录;GET /inbound/order/page 入库单分页查询;POST /inbound/order/save 保存入库单草稿;POST /inbound/order/submit/{id} 提交入库单;POST /outbound/order/apply 提交出库申请;POST /outbound/order/approval/handle 出库审批处理;GET /stock/list/page 库存分页查询;GET /report/dashboard 仪表盘数据。")

add_h2(doc, "6.5  安全设计")
for t in [
    "密码加密:BCrypt(强度 10),不可逆加密存储;",
    "Token 鉴权:Sa-Token 生成 UUID,存储于 Redis,有效期 8 小时;",
    "权限控制:@SaIgnore 标注公开接口(登录、验证码),其余接口由 Sa-Token 拦截器统一校验登录态;前端 Vue Router beforeEach + Pinia 按角色动态过滤路由,实现前后端双重权限控制;",
    "操作日志:自定义 @Log 注解 + AOP 切面,记录模块/操作/方法/URL/IP/参数/响应/耗时/状态/错误信息,写入 sys_operation_log 表;",
        "异常处理:GlobalExceptionHandler 统一捕获业务异常(参数错误、库存不足、状态非法、用户禁用等),返回标准 Result<T> 错误响应。",
]
add_list_item(doc, t)


# ============================================
# 第 5 章
# ============================================
add_h1(doc, "6  系统实现")

add_h2(doc, "7.1  开发环境")
add_three_line_table(doc,
    ["工具", "版本", "用途"],
    [
        ["JDK", "1.8", "运行环境"],
        ["Maven", "3.8+", "项目构建"],
        ["Node.js", "18+", "前端运行环境"],
        ["MySQL", "8.0.33", "数据库"],
        ["Redis", "7.x", "缓存/Session"],
        ["IDEA", "2023", "后端开发"],
        ["VS Code", "最新", "前端开发"]
    ],
    col_widths_cm=[3, 3, 9.5],
    caption="表 5.1  开发环境与工具版本"
)

add_h2(doc, "7.2  公共模块实现")
add_h3(doc, "6.2.1  统一返回 Result")
add_para(doc, "Result<T> 统一封装返回结果,包含 code、message、data、timestamp 四个字段,提供 ok()/ok(T data)/fail(String msg) 三个静态工厂方法,所有 Controller 均返回该类型。")
add_code_block(doc, [
    "public class Result<T> {",
    "    private Integer code;",
    "    private String message;",
    "    private T data;",
    "    private Long timestamp;",
    "    public static <T> Result<T> ok() { ... }",
    "    public static <T> Result<T> ok(T data) { ... }",
    "    public static <T> Result<T> fail(String msg) { ... }",
    "}",
])

add_h3(doc, "7.2.2  全局异常处理")
add_para(doc, "GlobalExceptionHandler 统一处理:BizException(业务异常,返回业务错误码)、NotLoginException(未登录,返回 401)、NotPermissionException(无权限,返回 403)、MethodArgumentNotValidException(参数校验失败,返回 400),所有异常最终包装为 Result.fail 统一返回。")

add_h3(doc, "7.2.3  操作日志 AOP")
add_para(doc, "通过 @Around(\"@annotation(com.wms.framework.annotation.Log)\") 切面拦截,记录方法执行耗时,异常情况下记录错误信息,最终异步写入 sys_operation_log 表。关键代码片段如下:")
add_code_block(doc, [
    '@Around("@annotation(com.wms.framework.annotation.Log)")',
    "public Object around(ProceedingJoinPoint pjp) {",
    "    long t = System.currentTimeMillis();",
    "    try { return pjp.proceed(); }",
    "    finally {",
    "        long cost = System.currentTimeMillis() - t;",
    "        saveLog(pjp, cost);  // 异步写入 sys_operation_log",
    "    }",
    "}",
])

add_h2(doc, "7.3  用户权限模块")
add_para(doc, "使用 Sa-Token 实现登录认证、Token 校验、角色权限控制。登录流程:验证 BCrypt 密码 → StpUtil.login(userId) → 写入 Redis → 返回 Token,后续请求在 Header 中携带 Authorization 即可访问鉴权接口。")

add_h2(doc, "7.4  入库管理(核心功能)")
add_h3(doc, "6.4.1  状态机")
add_para(doc, "入库单状态:DRAFT → PENDING → APPROVED → EXECUTING → FINISHED,或 REJECTED / CANCELED。状态机由 Service 层统一校验,确保状态流转的合法性。具体状态及可执行操作如下表所示:")
add_three_line_table(doc,
    ["状态", "说明", "可执行操作"],
    [
        ["DRAFT", "草稿", "编辑、提交、作废"],
        ["PENDING", "待审核", "审核(通过/驳回)"],
        ["APPROVED", "已审核", "执行"],
        ["EXECUTING", "执行中", "完成"],
        ["FINISHED", "已完成", "—"],
        ["REJECTED", "已驳回", "—"],
        ["CANCELED", "已作废", "—"]
    ],
    col_widths_cm=[3, 3, 9.5],
    caption="表 5.2  入库单状态机"
)

add_h3(doc, "7.4.2  入库核心事务")
add_para(doc, "入库执行的并发安全由 SELECT ... FOR UPDATE 行级锁 + Spring 事务保证,核心代码如下:")
add_code_block(doc, [
    "@Transactional(rollbackFor = Exception.class)",
    "public void executeInbound(String orderNo,",
    "        List<StockChangeItem> changes, Long operatorId) {",
    "    for (StockChangeItem ch : changes) {",
    "        // SELECT ... FOR UPDATE 行级锁,避免并发覆盖",
    "        Stock s = baseMapper.selectForUpdate(",
    "            ch.getGoodsId(), ch.getLocationId(), ch.getBatchNo());",
    "        int before = 0, after = 0;",
    "        if (s == null) {",
    "            // 新增库存记录",
    "            s = new Stock();",
    "            s.setGoodsId(ch.getGoodsId());",
    "            s.setLocationId(ch.getLocationId());",
    "            s.setBatchNo(ch.getBatchNo());",
    "            baseMapper.insert(s);",
    "        }",
    "        before = s.getQuantity();",
    "        after = before + ch.getQty();",
    "        s.setQuantity(after);",
    "        s.setAvailableQty(after - s.getLockedQty());",
    "        s.setLastInTime(LocalDateTime.now());",
    "        baseMapper.updateById(s);",
    "        // 写库存流水 (略)",
    "    }",
    "}",
])
add_para(doc, "关键技术点:SELECT ... FOR UPDATE 行级锁,在事务内锁定库存记录,避免并发修改覆盖;写库存 + 写流水在同一事务内,保证原子性;流水表记录业务单号,便于追溯与审计。")

add_h2(doc, "7.5  出库管理(两级审批)")
add_h3(doc, "6.5.1  状态机")
add_para(doc, "出库单状态:APPLY → APPROVING → APPROVED → SHIPPED → FINISHED,或 REJECTED / CANCELED(拣货与发货合并在 ship 一步完成,行级锁扣减库存)。")

add_h3(doc, "7.5.2  两级审批流程")
add_para(doc, "step=1 (部门负责人):APPLY → APPROVING(通过) / REJECTED(驳回);step=2 (仓库管理员):APPROVING → APPROVED(通过) / REJECTED(驳回)。每次审批写入 wms_outbound_approval 表留痕,便于审计追溯。")

add_h3(doc, "7.5.3  出库执行")
add_para(doc, "executeOutbound 方法实现出库执行:对每行明细加行级锁后,校验库存是否充足(可用数量 < 扣减数量时抛出 STOCK_NOT_ENOUGH 异常),更新库存数量与可用数量,同步写入 OUTBOUND 类型流水。")

add_h2(doc, "7.6  库存与盘点")
add_h3(doc, "6.6.1  实时库存")
add_para(doc, "联合唯一索引 (goods_id, location_id, batch_no) 保证同一商品在同一库位同一批次只能存在一条记录,杜绝重复库存。available_qty = quantity - locked_qty 实时计算可用量,支持预占与释放。")

add_h3(doc, "7.6.2  盘点流程")
add_para(doc, "创建盘点单 → 自动抓取当前库存(system_qty)→ 录入实盘数(actual_qty)→ 计算差异(diff_qty = actual - system)→ 确认调整 → 写盘点调整流水。")

add_h2(doc, "7.7  报表与预警")
add_h3(doc, "6.7.1  仪表盘")
add_para(doc, "使用 ECharts 5 折线图分别展示近 7 天入库与出库趋势;数字卡片展示总库存量、今日入库/出库、预警数量;下方列表展示具体的预警通知。")

add_h3(doc, "7.7.2  库存预警")
add_inline_para(doc, [
    ("通过", {'name': F_SONG, 'size': 12}),
    ("@Scheduled(cron = \"0 0 * * * ?\") ", {'name': 'Consolas', 'size': 11}),
    ("注解实现定时任务,每小时执行一次,扫描低于安全库存的商品与接近保质期的商品,将预警信息写入 wms_notification 表。", {'name': F_SONG, 'size': 12})
])

add_h2(doc, "7.8  前端实现")
add_para(doc, "前端采用 Vue 3 + Vite + Element Plus + Pinia + Vue Router 技术栈:状态管理采用 Pinia 存储用户信息、Token、路由权限;路由控制根据用户角色动态生成菜单和路由;HTTP 请求采用 Axios 封装统一请求拦截器,自动携带 Token;页面组件共 23+ 业务页面,包括列表页、表单页、详情页。")

add_page_break(doc)

# ============================================
# 第 6 章
# ============================================
add_h1(doc, "7  系统测试")

add_h2(doc, "8.1  测试环境")
add_para(doc, "测试环境包括:Windows 11 操作系统、MySQL 8.0.33 数据库、Redis 7.0 缓存、Chrome / Edge / Firefox 主流浏览器。")

add_h2(doc, "8.2  功能测试用例")
add_three_line_table(doc,
    ["模块", "测试项", "预期结果", "实际结果", "结论"],
    [
        ["登录", "admin/123456", "登录成功", "成功", "通过"],
        ["登录", "错误密码", "提示密码错误", "提示", "通过"],
        ["用户管理", "新增用户", "用户列表显示", "显示", "通过"],
        ["商品管理", "新增商品", "列表显示", "显示", "通过"],
        ["库位管理", "新增库位", "库位树显示", "显示", "通过"],
        ["入库", "仓管创建草稿", "状态=DRAFT", "DRAFT", "通过"],
        ["入库", "提交入库单", "状态=PENDING", "PENDING", "通过"],
        ["入库", "审核通过", "状态=APPROVED", "APPROVED", "通过"],
        ["入库", "执行完成", "库存+", "正确", "通过"],
        ["出库", "员工申请出库", "状态=APPLY", "APPLY", "通过"],
        ["出库", "部门负责人审核", "状态=APPROVING", "APPROVING", "通过"],
        ["出库", "仓管审核通过", "状态=APPROVED", "APPROVED", "通过"],
        ["出库", "发货出库", "库存-,状态=SHIPPED", "正确", "通过"],
        ["库存", "库存查询", "显示库存数据", "显示", "通过"],
        ["预警", "低库存扫描", "生成预警通知", "是", "通过"],
        ["权限", "普通员工登录后访问 /system/user 页面", "路由守卫拦截,菜单不显示", "已拦截", "通过"],
    ],
    col_widths_cm=[2, 4.5, 3.5, 3, 1.5],
    caption="表 6.1  系统功能测试用例与结果"
)

add_h2(doc, "8.3  性能测试")
add_para(doc, "系统采用 Spring Boot + Sa-Token + MyBatis-Plus + Redis 技术栈,结合行级锁与连接池调优,接口响应理论上能够满足本系统预期的并发与响应要求。正式上线前应使用 JMeter 或 wrk 等工具对核心接口(登录、库存查询、仪表盘、入出库提交等)进行压测,记录实际平均响应时间与 P99 延迟,本节具体压测数据待补。")

add_h2(doc, "8.4  兼容性测试")
add_para(doc, "前端基于 Vue 3 + Element Plus 构建,组件库基于现代浏览器标准,理论上支持 Chrome、Edge、Firefox 等主流浏览器的最近 2 个主版本。由于 Element Plus 不再支持 IE 浏览器,本系统明确不兼容 IE 11 及以下版本。本节具体的浏览器版本与分辨率兼容性测试数据待补。")

add_page_break(doc)

# ============================================
# 第 7 章
# ============================================
add_h1(doc, "8  总结与展望")

add_h2(doc, "8.1  工作总结")
add_para(doc, "本文设计并实现了一套基于 Spring Boot + Vue 3 的电商仓储物资管理系统,主要成果包括:")
for i, t in enumerate([
    "完成系统需求分析与设计:调研电商仓储管理需求,完成功能需求与非功能需求分析,设计系统总体架构与数据库结构;",
    "完成 21 张业务数据库表设计:实现完整 E-R 模型,包含用户权限、主数据、入库、出库、库存、盘点、通知等业务表;",
    "实现 4 类角色 RBAC 权限管理:基于 Sa-Token 实现登录认证、Token 校验、角色权限控制;",
    "实现入库单/出库单状态机:入库 7 状态、出库 8 状态,通过 Service 层统一校验状态合法性;",
    "实现库存并发安全机制:通过 SELECT ... FOR UPDATE 行级锁与事务,保证库存数据一致性;",
    "实现库存预警定时任务:每小时扫描低库存与临期商品,自动生成预警通知;",
    "实现 ECharts 数据可视化:仪表盘展示总库存、入出库趋势、预警数量等关键指标;",
    "完成 8 份毕设文档:开题报告、毕业论文、答辩 PPT、答辩稿、数据库设计文档、接口文档、用户手册、演示脚本。",
], 1):
    add_list_item(doc, t, idx=i)

add_h2(doc, "8.2  不足与展望")
add_para(doc, "虽然系统已满足基本功能需求,但仍有以下可优化方向:")
for i, t in enumerate([
    "微服务化:将库存服务、单据服务、报表服务拆分为独立微服务,使用 Spring Cloud 或 Dubbo,实现更高水平的扩展性;",
    "移动端支持:开发移动端 H5 页面或微信小程序,支持扫码作业,提升现场作业效率;",
    "AI 补货建议:基于历史销售数据,使用机器学习算法预测未来需求,自动生成采购建议;",
    "IoT 集成:对接电子标签、RFID、AGV 等智能仓储设备,实现自动化作业;",
    "多租户 SaaS:支持多企业共用一套系统,通过数据隔离实现多租户支持。",
], 1):
    add_list_item(doc, t, idx=i)

add_page_break(doc)

# ============================================
# 参考文献
# ============================================
add_title_center(doc, "参 考 文 献", size=16, name=F_HEI)
refs = [
    "[1] 汪云飞. Java EE 开发的颠覆者:Spring Boot 实战[M]. 北京:电子工业出版社,2016.",
    "[2] 尤雨溪. Vue.js 设计与实现[M]. 北京:人民邮电出版社,2022.",
    "[3] 杨开振. 深入浅出 Spring Boot 2.x[M]. 北京:人民邮电出版社,2019.",
    "[4] Craig Walls. Spring 实战(第 4 版)[M]. 张卫滨译. 北京:人民邮电出版社,2016.",
    "[5] 廖雪峰. SQL 入门教程[EB/OL]. https://www.liaoxuefeng.com,2023.",
    "[6] 腾讯云. Element Plus 官方文档[EB/OL]. https://element-plus.org,2024.",
    "[7] 百度. Echarts 数据可视化[EB/OL]. https://echarts.apache.org,2024.",
    "[8] 阿里巴巴. MyBatis-Plus 官方文档[EB/OL]. https://baomidou.com,2024.",
    "[9] 一只野生程序猿. Sa-Token 轻量级 Java 权限认证框架[EB/OL]. https://sa-token.dev33.cn,2024.",
    "[10] 杨恩雄. 高性能 MySQL(第 3 版)[M]. 宁海元译. 北京:电子工业出版社,2013.",
    "[11] 陈皓. 分布式系统的事务处理[J]. 程序员,2014(05):12-18.",
    "[12] Martin Fowler. Patterns of Enterprise Application Architecture[M]. Addison-Wesley,2002.",
    "[13] Eric Evans. 领域驱动设计[M]. 赵俐译. 北京:人民邮电出版社,2010.",
    "[14] 崔鹏飞. 仓储管理系统设计与实现[D]. 北京:北京邮电大学,2022.",
    "[15] 王伟. 基于 Spring Boot 的电商仓储管理系统的研究[D]. 上海:东华大学,2023.",
    "[16] ZDNet. Warehouse Management System Market Report[R]. 2023.",
    "[17] Gartner. Magic Quadrant for WMS[R]. 2024.",
]
for r in refs:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    run = p.add_run(r)
    set_cn_font(run, name=F_SONG, size=10.5)

add_page_break(doc)

# ============================================
# 致谢
# ============================================
add_title_center(doc, "致  谢", size=16, name=F_HEI)
add_para(doc, "时光荏苒,四年的本科学习即将结束。本文是在任占广老师的悉心指导下完成的,从选题到方案设计,从编码实现到论文撰写,老师都给予了耐心的指导与帮助。在此向老师致以诚挚的谢意。")
add_para(doc, "同时感谢实验室的同学们,在系统设计与实现过程中,我们一起讨论方案、解决 Bug,共同度过了难忘的时光。")
add_para(doc, "最后感谢我的家人,他们的支持与鼓励是我前行的最大动力。")

add_page_break(doc)

# ============================================
# 附录
# ============================================
add_title_center(doc, "附  录", size=16, name=F_HEI)
add_para(doc, "附录 A:数据库脚本(详见 wms-db/02_create_tables.sql)", indent=False)
add_para(doc, "附录 B:关键代码(详见 wms-backend/src/main/java/com/wms/)", indent=False)
add_para(doc, "附录 C:用户手册(详见 docs/07-用户手册.md)", indent=False)
add_para(doc, "附录 D:演示截图与脚本(详见 docs/screenshots/、08-演示脚本.md)", indent=False)

# ============================================
# 页眉页脚
# ============================================
section.different_first_page_header_footer = True
# 首页(封面)无header/footer
fph = section.first_page_header
fph.is_linked_to_previous = False
if fph.paragraphs:
    for p in fph.paragraphs:
        for r in p.runs:
            r.text = ""
fpf = section.first_page_footer
fpf.is_linked_to_previous = False
if fpf.paragraphs:
    for p in fpf.paragraphs:
        for r in p.runs:
            r.text = ""

# 其他页的页眉:论文题目
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
hp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
hp.paragraph_format.first_line_indent = Cm(0)
# 文字
if not hp.runs:
    run = hp.add_run(TITLE_LINE1 + TITLE_LINE2)
else:
    run = hp.runs[0]
    run.text = TITLE_LINE1 + TITLE_LINE2
set_cn_font(run, name=F_SONG, size=9)
# 下划线
pPr = hp._element.get_or_add_pPr()
# 清除已有pBdr
for old in pPr.findall(qn('w:pBdr')):
    pPr.remove(old)
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '6')
bottom.set(qn('w:space'), '1')
bottom.set(qn('w:color'), 'auto')
pBdr.append(bottom)
pPr.append(pBdr)

# 页脚: - N -
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
fp.paragraph_format.first_line_indent = Cm(0)
if not fp.runs:
    run = fp.add_run("- ")
    set_cn_font(run, name=F_TIMES, size=9)
    # PAGE field
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run2 = fp.add_run()
    set_cn_font(run2, name=F_TIMES, size=9)
    run2._element.append(fld_begin)
    run2._element.append(instr)
    run2._element.append(fld_sep)
    run3 = fp.add_run("1")
    set_cn_font(run3, name=F_TIMES, size=9)
    run4 = fp.add_run()
    set_cn_font(run4, name=F_TIMES, size=9)
    run4._element.append(fld_end)
    run5 = fp.add_run(" -")
    set_cn_font(run5, name=F_TIMES, size=9)

doc.save(OUT)
print(f"OK: {OUT}")
print(f"Size: {os.path.getsize(OUT)} bytes")
